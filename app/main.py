from fastapi import FastAPI, Request, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from pathlib import Path
from typing import Dict
import unicodedata
import re
import subprocess
import uuid

from docx import Document  # python-docx

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"
SAIDA_DIR = BASE_DIR / "saida"

SAIDA_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Jul.IA – Automação de Procuração e Contrato de Consignado")

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


# --------------------- HELPERS ---------------------


def parse_bloco_dados(bloco: str) -> Dict[str, str]:
    """
    Lê bloco no formato:
    Nome completo: ...
    Nacionalidade: ...
    etc.
    Retorna dicionário com chaves em minúsculo.
    """
    dados: Dict[str, str] = {}
    for linha in bloco.splitlines():
        linha = linha.strip()
        if not linha or ":" not in linha:
            continue
        rotulo, valor = linha.split(":", 1)
        rotulo = rotulo.strip().lower()
        valor = valor.strip()
        dados[rotulo] = valor
    return dados


def extrair_nome_completo(dados: Dict[str, str]) -> str:
    """
    Tenta encontrar o nome completo em várias chaves.
    """
    possiveis = [
        "nome completo",
        "nome",
        "cliente",
        "autor",
    ]
    for chave in possiveis:
        if chave in dados and dados[chave].strip():
            return dados[chave].strip()
    return "Cliente Autor"


def slug_nome_sobrenome(nome_completo: str) -> str:
    """
    A partir de 'Juliano Garbuggio da Silva' -> 'Juliano_Garbuggio'
    Remove acentos e caracteres esquisitos.
    """
    # Normaliza/remover acentos
    nome_norm = unicodedata.normalize("NFKD", nome_completo)
    nome_ascii = "".join(c for c in nome_norm if not unicodedata.combining(c))
    # Mantém só letras, números e espaços
    nome_limpo = re.sub(r"[^A-Za-z0-9 ]+", "", nome_ascii).strip()
    if not nome_limpo:
        return "Cliente_Autor"

    partes = nome_limpo.split()
    if len(partes) == 1:
        primeiro = partes[0].capitalize()
        sobrenome = "Autor"
    else:
        primeiro = partes[0].capitalize()
        sobrenome = partes[-1].capitalize()
    return f"{primeiro}_{sobrenome}"


def montar_nome_arquivo(nome_completo: str, ext: str) -> Path:
    """
    Gera:
    02_Procuracao_Kit_Consignado_Nome_Sobrenome_Autor.ext
    """
    slug = slug_nome_sobrenome(nome_completo)
    base_nome = f"02_Procuracao_Kit_Consignado_{slug}_Autor"
    return SAIDA_DIR / f"{base_nome}.{ext.lstrip('.')}"


def gerar_docx_a_partir_dos_dados(dados: Dict[str, str], caminho_saida: Path) -> None:
    """
    Gera um DOCX simples só para garantir funcionamento.
    Se você tiver um template próprio, pode adaptar este trecho
    para usar docxtpl ou carregar seu modelo.
    """
    doc = Document()
    doc.add_heading("Procuração + Contrato de Empréstimo Consignado", level=1)

    nome = extrair_nome_completo(dados)
    doc.add_paragraph(f"Cliente: {nome}")

    doc.add_heading("Dados do cliente", level=2)
    for chave, valor in dados.items():
        linha = f"{chave.capitalize()}: {valor}"
        doc.add_paragraph(linha)

    doc.add_paragraph("")
    doc.add_paragraph("Documento gerado automaticamente por Jul.IA.")

    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(caminho_saida))


def converter_docx_para_pdf(caminho_docx: Path, caminho_pdf: Path) -> None:
    """
    Usa LibreOffice em modo headless para converter DOCX -> PDF.
    É o padrão que funciona bem no Railway quando o LibreOffice
    está instalado na imagem Docker.
    """
    comando = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(caminho_pdf.parent),
        str(caminho_docx),
    ]
    subprocess.run(comando, check=True)
    # LibreOffice gera o PDF com o mesmo nome do DOCX
    # na pasta destino; garantimos que o nome final seja o esperado.
    gerado = caminho_pdf.parent / (caminho_docx.stem + ".pdf")
    if gerado != caminho_pdf and gerado.exists():
        gerado.rename(caminho_pdf)


# --------------------- ROTAS WEB ---------------------


@app.get("/", response_class=HTMLResponse)
async def dashboard(request: Request):
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "titulo_pagina": "Jul.IA – Automação de Procuração e Contrato de Consignado",
        },
    )


@app.get("/docx", response_class=HTMLResponse)
async def pagina_docx(request: Request):
    return templates.TemplateResponse(
        "docx.html",
        {
            "request": request,
            "titulo_pagina": "Jul.IA – Automação de Procuração e Contrato de Consignado",
        },
    )


@app.get("/pdf", response_class=HTMLResponse)
async def pagina_pdf(request: Request):
    return templates.TemplateResponse(
        "pdf.html",
        {
            "request": request,
            "titulo_pagina": "Jul.IA – Automação de Procuração e Contrato de Consignado",
        },
    )


@app.post("/gerar-docx")
async def gerar_docx(bloco_dados: str = Form(...)):
    dados = parse_bloco_dados(bloco_dados)
    nome_completo = extrair_nome_completo(dados)
    caminho_docx = montar_nome_arquivo(nome_completo, "docx")

    gerar_docx_a_partir_dos_dados(dados, caminho_docx)

    return FileResponse(
        path=str(caminho_docx),
        filename=caminho_docx.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/gerar-pdf")
async def gerar_pdf(bloco_dados: str = Form(...)):
    dados = parse_bloco_dados(bloco_dados)
    nome_completo = extrair_nome_completo(dados)
    caminho_docx = montar_nome_arquivo(nome_completo, "docx")
    caminho_pdf = montar_nome_arquivo(nome_completo, "pdf")

    gerar_docx_a_partir_dos_dados(dados, caminho_docx)
    converter_docx_para_pdf(caminho_docx, caminho_pdf)

    return FileResponse(
        path=str(caminho_pdf),
        filename=caminho_pdf.name,
        media_type="application/pdf",
    )


# Para rodar local:
# uvicorn app.main:app --reload --host 0.0.0.0 --port 8011
